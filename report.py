import pandas as pd
import os
import sys
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tkinter as tk
from tkinter import filedialog,messagebox
from docx.enum.text import WD_UNDERLINE
import subprocess
import sys


def install_package(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

# def check_and_install_packages(required_packages):
#     for package in required_packages:
#         try:
#             __import__(package)
#         except ImportError:
#             print(f"{package} is not installed. Installing...")
#             install_package(package)
#             print(f"{package} installed successfully.")

# required_packages = ['pandas', 'openpyxl', 'python-docx', 'docx', 'tkinter']

# 在脚本开始时检查并安装缺失的库

Excel_path = ""
Output_path = ""
Word_template_path = ""
date = ""
period_number = ""

def preprocess_str(row):
    # 将每段内容按换行符分割
    paragraphs = row.split('\n')
    
    # 处理每段内容
    processed_paragraphs = []
    for paragraph in paragraphs:
        # 替换之前的空格为4个空格
        paragraph = paragraph.replace(' ', '    ')  # 4个空格
        # 在每段前添加4个空格
        processed_paragraphs.append('    ' + paragraph)
    
    # 将处理后的段落重新合并为一个字符串
    return '\n'.join(processed_paragraphs)


#读取Excel
def read_excel(file_path):
    return pd.read_excel(file_path)

#写入Excel
def write_excel(data,output_path):
    data.to_excel(output_path,index = False)

#创建文件夹
def create_folder(folder_path):
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

#把题目分割：AI applications to landslide risk assessment_Filippo Catani_20241011_15:30-17:00_测绘馆_431
def extract(title):
    parts = title.split('_')
    if len(parts) < 6:
        raise ValueError("Title格式不正确，应包含至少6个部分")
    result = {
    'topic' : parts[0],
    'author' : parts[1],
    'date' : parts[2],
    'time' : parts[3],
    'location' : parts[4],
    'period_number' : parts[5]
    }
    return result
    
#Excel的预处理，把问卷的内容写到新的Excel中，方便接下来的处理。

def PreProcess(Ori_file_path,output_path):
    #读取原始数据表格
    df = read_excel(Ori_file_path)
    #表头选取
    columns_to_write = ['4、报告名称：', '1、姓名', '2、学号', '3、年级', '5、主要体会']
    df_select = df[columns_to_write]
    df_select.columns = ['Title', 'Name', 'ID', 'Grade', 'Content']
    #这里是一个总的excel，汇总了本次的所有报告
    Sum_path = output_path + "\\Sum.xlsx"
    #Sum_path="D:\\Code\\py_projects\\Report \\Sum.xlsx"
    write_excel(df_select,Sum_path)
    print("success")
    #分类的操作在这里实现
    df_select = read_excel(Sum_path)
    #遍历sum，按年级分类，创建folder
    grades = df_select['Grade'].unique()
     #其实可以嵌套进主函数，让这个遍历之后直接把文档扔文件夹里
    for grade in grades:
        folder_path_grades = output_path + '\\reports\\' + f'{grade}'
        df_grade = df_select[df_select["Grade"]==grade]
        path = folder_path_grades + '\\' + f'{grade}.xlsx'
        create_folder(folder_path_grades)
        write_excel(df_grade,path)     
    print("表格已完成预处理")
#PreProcess("D:\\Code\\py_projects\\Report\\286544793_按文本_科技报告提交示例_7_2.xlsx","D:\\Code\\py_projects\\Report")

def pad_string(n,input_str):
    # 计算需要添加的空格数
    if len(input_str) < 16:
        padding = (n - len(input_str)) // 2
    # 在字符串前后添加空格
        padded_str = ' ' * padding + input_str + ' ' * padding
        return padded_str
    else:
        return input_str

def insert_data_to_word(date,period_number,template_path,output_path):
    #日期
    #date = 20241020#加在问卷里，问卷题目设置的时候就注明date，然后读取即可
    #从Excel里导入

    #期号
    #period_number = 422#加在问卷里，问卷题目设置的时候就注明date，然后读取即可

    #从Excel里导入
    sum_path = output_path + "\\Sum.xlsx"
    data = read_excel(sum_path)
    #data换成路径
    #总文档文件夹的路径，后面加的是日期。

    folder_word_path = output_path + f"\\{date}"
    grades = data['Grade'].unique()
    print
    #遍历年级
    for grade in grades:
        folder_word_path_grades = folder_word_path + '\\reports_word_grades\\' + f'{grade}'
        create_folder(folder_word_path_grades)
        #excel的位置：
        grade_excel_path = output_path + '\\reports\\' + f'{grade}' + '\\' + f'{grade}.xlsx'
        data = read_excel(grade_excel_path)
        for row in data.itertuples():#这里应该是Excel的位置
            doc = Document(template_path)
            title = str(row.Title)
            try:
                Title_dic = extract(title)
                Title_dic["name"] = str(row.Name)
                Title_dic["ID"] = str(row.ID)
                date = str(Title_dic['date'])
                #20241021
                year = date[0] + date[1] +date[2] +date[3]
                month = date[4] + date[5]
                day = date[6] +date[7]
                Date = year +"年" + month + "月" + day + "日"
                period_number = Title_dic['period_number']
            except ValueError as e:
                print(f"错误：{e} - Title: {title}")
                continue
                paragraph_format.first_line_indent = indent_in
            for paragraph in doc.paragraphs:
                # paragraph_format = paragraph.paragraph_format
                # paragraph_format.first_line_indent = Pt(12)
                runs = paragraph.runs
                for run in runs:
                    if '{topic}' in run.text:
                        run.text = run.text.replace("{topic}", Title_dic['topic'])
                        run.font.underline = WD_UNDERLINE.SINGLE
                        run.font.size = Pt(14)
                    elif '{author}' in run.text:
                        run.text = run.text.replace("{author}", pad_string(18,Title_dic['author']))
                        run.font.underline = WD_UNDERLINE.SINGLE
                        run.font.size = Pt(12)
                    elif '{location}' in run.text:
                        run.text = run.text.replace("{location}",pad_string(12,Title_dic['location']))
                        run.font.underline = WD_UNDERLINE.SINGLE
                        run.font.size = Pt(12)
                    elif '{date}' in run.text:
                        run.text = run.text.replace("{date}",pad_string(16,Date))
                        run.font.underline = WD_UNDERLINE.SINGLE
                        run.font.size = Pt(12)
                    elif '{time}' in run.text:
                        run.text = run.text.replace("{time}",pad_string(18,f"{Title_dic['time']}"))
                        run.font.underline = WD_UNDERLINE.SINGLE
                        run.font.size = Pt(12)
                    elif '{name}' in run.text:
                        run.text = run.text.replace("{name}",pad_string(18,Title_dic['name']))
                        run.font.underline = WD_UNDERLINE.SINGLE
                        run.font.size = Pt(12)
                    elif '{ID}' in run.text:
                        run.text = run.text.replace("{ID}",pad_string(16,f"{Title_dic['ID']}"))
                        run.font.underline = WD_UNDERLINE.SINGLE
                        run.font.size = Pt(12)
                    elif '{grade}' in run.text:
                        run.text = run.text.replace("{grade}",pad_string(14,f"{row.Grade}"))
                        run.font.underline = WD_UNDERLINE.SINGLE
                        run.font.size = Pt(11)
                    elif '{content}' in paragraph.text:
                        #original_runs = paragraph.runs
                        # paragraph_format = paragraph.paragraph_format
                        # paragraph_format.first_line_indent = Pt(12)
                        paragraph.clear()
                        run = paragraph.add_run(preprocess_str(str(row.Content)))
                        run.font.size = Pt(12)
                    else:
                        original_runs = paragraph.runs
            file_name = f"{date}-{period_number} {row.ID} {row.Name} 科技报告听课证明.docx"
            #报告的生成的路径，这里用
            output_word_path = folder_word_path_grades + "\\"  + file_name
            #防止重复
            if not os.path.exists(output_word_path):
                doc.save(output_word_path)
            #output_path2 = "D:\\Code\\py_projects\\Report\\20241017\\20241017-422 2250450 郑盛国 科技报告听课证明.docx"
            #其实可以嵌套进主函数，让这个遍历之后直接把文档扔文件夹里
#insert_data_to_word(20241020,422,"D:\\Code\\py_projects\\Report\\Report2.docx","D:\\Code\\py_projects\\Report")

# 主函数
# def main():
#     #日期
#     date = 20241027
#     #处理序号
#     period_number = 1
#     #问卷信息
#     Ori_path = "D:\\Code\\py_projects\\Report\\286544793_按文本_科技报告提交示例_8_1.xlsx"  # 替换为你的Excel文件路径
#     #输出文件夹
#     output_path = "D:\\Code\\py_projects\\Report"  # 替换为你想要保存的新Excel文件路径
#     PreProcess(Ori_path,output_path)
#     # Word模板路径(不同专业加一行文件处理的内容：)
#     word_template_path = "D:\Code\py_projects\Report\Report2.docx"
#     # 读取Excel数据
#     insert_data_to_word(date,period_number,word_template_path,output_path)

# if __name__ == "__main__":
#     main()


def choose_excel_path():
    global Excel_path
    Excel_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls")])
    if Excel_path:
        excel_path_label.config(text=Excel_path)
    else:
        excel_path_label.config(text="Excel的路径")


def choose_output_path():
    global Output_path
    Output_path = filedialog.askdirectory()
    if Output_path:
        output_path_label.config(text=Output_path)
    else:
        output_path_label.config(text="输出路径")

def choose_word_template_path():
    global Word_template_path
    Word_template_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx *.doc")])
    if Word_template_path:
        word_template_path_label.config(text=Word_template_path)
    else:
        word_template_path_label.config(text="Word模版的路径")

def refresh_paths():
    global Excel_path, Output_path, Word_template_path
    Excel_path, Output_path, Word_template_path = "", "", ""
    excel_path_label.config(text="Excel的路径")
    output_path_label.config(text="输出路径")
    word_template_path_label.config(text="Word模版的路径")



def start_main():
    # # 这里应该是你的main函数的代码
    # check_and_install_packages(required_packages)
    global date, period_number
    # 获取输入框的值
    date = date_entry.get() if date_entry.get() else "日期"
    period_number = period_number_entry.get() if period_number_entry.get() else "序号"
    # #问卷信息
    # #Excel_path = "D:\\Code\\py_projects\\Report\\286544793_按文本_科技报告提交示例_8_1.xlsx"  # 替换为你的Excel文件路径
    # #输出文件夹
    # #output_path = "D:\\Code\\py_projects\\Report"  # 替换为你想要保存的新Excel文件路径
    # Excel_path = choose_excel_path()
    messagebox.showinfo("信息", "开始执行main函数")
    PreProcess(Excel_path,Output_path)
    # Word模板路径(不同专业加一行文件处理的内容：)
    # word_template_path = "D:\Code\py_projects\Report\Report2.docx"
    # 读取Excel数据
    insert_data_to_word(date,period_number,Word_template_path,Output_path)
    messagebox.showinfo("信息", f"执行完毕，输出路径为：{Output_path}")

# 初始化窗口
root = tk.Tk()
root.title("路径选择器")

# Excel路径选择
excel_path_label = tk.Label(root, text="Excel的路径", width=50)
excel_path_label.grid(row=0, column=1, padx=10, pady=10)

excel_path_button = tk.Button(root, text="选择Excel路径", command=choose_excel_path)
excel_path_button.grid(row=0, column=2, padx=10, pady=10)

# 输出路径选择
output_path_label = tk.Label(root, text="输出路径", width=50)
output_path_label.grid(row=1, column=1, padx=10, pady=10)

output_path_button = tk.Button(root, text="选择输出路径", command=choose_output_path)
output_path_button.grid(row=1, column=2, padx=10, pady=10)

# Word模板路径选择
word_template_path_label = tk.Label(root, text="Word模版的路径", width=50)
word_template_path_label.grid(row=2, column=1, padx=10, pady=10)

word_template_path_button = tk.Button(root, text="选择Word模板路径", command=choose_word_template_path)
word_template_path_button.grid(row=2, column=2, padx=10, pady=10)

# 日期输入框
date_label = tk.Label(root, text="日期")
date_label.grid(row=3, column=0, padx=10, pady=10)

date_entry = tk.Entry(root, width=50)
date_entry.grid(row=3, column=1, padx=10, pady=10, sticky="ew")
date_entry.insert(0, "日期")
date_entry.bind("<Key>", lambda e: (e.widget.delete(0, tk.END)) if e.widget.get() == "日期" else None)
date_entry.bind("<KeyRelease>", lambda e: (e.widget.insert(0, "日期")) if not e.widget.get() else None)

# 序号输入框
period_number_label = tk.Label(root, text="序号")
period_number_label.grid(row=4, column=0, padx=10, pady=10)

period_number_entry = tk.Entry(root, width=50)
period_number_entry.grid(row=4, column=1, padx=10, pady=10, sticky="ew")
period_number_entry.insert(0, "序号")
period_number_entry.bind("<Key>", lambda e: (e.widget.delete(0, tk.END)) if e.widget.get() == "序号" else None)
period_number_entry.bind("<KeyRelease>", lambda e: (e.widget.insert(0, "序号")) if not e.widget.get() else None)

# 刷新按钮
refresh_button = tk.Button(root, text="Refresh", command=refresh_paths)
refresh_button.grid(row=5, column=1, padx=10, pady=10)

# 开始按钮
start_button = tk.Button(root, text="Start", command=start_main)
start_button.grid(row=5, column=2, padx=10, pady=10)

# 主循环
root.mainloop()
